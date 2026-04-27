from collections import Counter
import sys

from docx import Document

sys.stdout.reconfigure(encoding="utf-8")

path = r"C:\Users\Даня\Desktop\ЛР-4 Оптимизация.docx"
doc = Document(path)

print("PARAGRAPHS:", len(doc.paragraphs))
print("STYLES:", Counter(par.style.name for par in doc.paragraphs))
print("--- SAMPLE ---")
for i, par in enumerate(doc.paragraphs[:160]):
    print(f"{i}: [{par.style.name}] {par.text!r}")

print("--- SECTIONS ---")
for i, section in enumerate(doc.sections):
    print(
        i,
        section.top_margin.cm,
        section.bottom_margin.cm,
        section.left_margin.cm,
        section.right_margin.cm,
    )

print("--- FORMATTING SAMPLE ---")
for idx in [0, 3, 13, 14, 15, 22, 40, 41, 43, 45, 65, 138, 149, 159]:
    par = doc.paragraphs[idx]
    runs = [
        (
            run.text,
            run.font.name,
            run.font.size.pt if run.font.size else None,
            run.bold,
            run.italic,
        )
        for run in par.runs
    ]
    fmt = par.paragraph_format
    print(
        idx,
        {
            "text": par.text,
            "style": par.style.name,
            "alignment": par.alignment,
            "first_line_indent": fmt.first_line_indent.pt if fmt.first_line_indent else None,
            "left_indent": fmt.left_indent.pt if fmt.left_indent else None,
            "space_before": fmt.space_before.pt if fmt.space_before else None,
            "space_after": fmt.space_after.pt if fmt.space_after else None,
            "line_spacing": fmt.line_spacing,
            "runs": runs,
        },
    )
