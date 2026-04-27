import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


sys.stdout.reconfigure(encoding="utf-8")


DESKTOP = Path(r"C:\Users\Даня\Desktop")


def set_run_font(run, bold=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)
    if bold is not None:
        run.bold = bold


def set_paragraph_line_spacing(paragraph, line_spacing=1.5):
    fmt = paragraph.paragraph_format
    fmt.line_spacing = line_spacing


def add_paragraph(
    doc,
    text="",
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_indent_cm=1.25,
    left_indent_cm=-1.18,
    bold_prefix=None,
    style=None,
):
    p = doc.add_paragraph(style=style)
    p.alignment = alignment
    fmt = p.paragraph_format
    if first_indent_cm is not None:
        fmt.first_line_indent = Cm(first_indent_cm)
    if left_indent_cm is not None:
        fmt.left_indent = Cm(left_indent_cm)
    set_paragraph_line_spacing(p)

    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, bold=True)
        second = p.add_run(text[len(bold_prefix):])
        set_run_font(second)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_center_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_line_spacing(p)
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    return p


def add_blank(doc, count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("")
        set_run_font(run)


def add_heading_like(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.left_indent = Cm(-1.18)
    set_paragraph_line_spacing(p)
    run = p.add_run(text)
    set_run_font(run, bold=True)
    return p


def add_code_line(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.left_indent = Cm(0)
    fmt.first_line_indent = Cm(0)
    set_paragraph_line_spacing(p)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)
    return p


def add_placeholder_figure(doc, number, caption):
    add_blank(doc, 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[МЕСТО ДЛЯ ВСТАВКИ ИЗОБРАЖЕНИЯ]")
    set_run_font(run)
    run.italic = True
    add_blank(doc, 1)
    add_center_paragraph(doc, f"Рисунок {number} – {caption}")


def force_update_fields(doc):
    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.175)
    section.right_margin = Cm(3.175)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(14)

    force_update_fields(doc)
    return doc


def add_cover_page(doc, lab_number, topic):
    add_center_paragraph(doc, "Министерство науки и высшего образования Российской Федерации")
    add_center_paragraph(doc, "Федеральное государственное бюджетное образовательное учреждение")
    add_center_paragraph(doc, "высшего образования")
    add_center_paragraph(doc, "«КУБАНСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ»", bold=True)
    add_center_paragraph(doc, "(ФГБОУ ВО «КубГУ»)")
    add_blank(doc, 1)
    add_center_paragraph(doc, "Факультет компьютерных технологий и прикладной математики")
    add_center_paragraph(doc, "Кафедра вычислительных технологий")
    add_blank(doc, 5)
    add_center_paragraph(doc, f"ЛАБОРАТОРНАЯ РАБОТА №{lab_number}", bold=True)
    add_center_paragraph(doc, "Дисциплина: Метод поисковой оптимизации", bold=True)
    add_center_paragraph(doc, f"Тема: «{topic}»")
    add_blank(doc, 6)

    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_line_spacing(p1)
    r1 = p1.add_run("Работу выполнил: __________________________________")
    set_run_font(r1)
    r2 = p1.add_run("Д.Е.Котов")
    set_run_font(r2)

    add_blank(doc, 4)
    add_center_paragraph(
        doc,
        "Направление подготовки: 02.03.02 Фундаментальная информатика и информационные технологии",
    )
    add_blank(doc, 4)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_line_spacing(p2)
    r1 = p2.add_run("Преподаватель:__________________________________")
    set_run_font(r1)
    r2 = p2.add_run("Е.А.Нигодин")
    set_run_font(r2)

    add_blank(doc, 7)
    add_center_paragraph(doc, "Краснодар 2026")
    doc.add_page_break()


def build_lab5():
    doc = setup_document()
    add_cover_page(doc, 5, "АЛГОРИТМ ПЧЕЛИНОГО РОЯ")

    add_paragraph(
        doc,
        "Цель: изучить принципы работы алгоритма пчелиного роя как метода глобальной оптимизации, "
        "реализовать программно механизмы разведки и локального поиска вокруг перспективных участков, "
        "а также применить алгоритм для минимизации тестовых функций двух переменных.",
        bold_prefix="Цель:",
    )
    add_blank(doc, 1)

    add_heading_like(doc, "ТЕОРИТИЧЕСКАЯ ЧАСТЬ")
    add_blank(doc, 1)
    add_paragraph(
        doc,
        "Алгоритм пчелиного роя относится к популяционным стохастическим методам оптимизации. "
        "Он моделирует поведение колонии пчел, в которой часть агентов выполняет роль разведчиков, "
        "а часть направляется в наиболее перспективные зоны пространства поиска. Такой подход сочетает "
        "глобальное исследование области и локальное уточнение найденных хороших решений."
    )
    add_paragraph(
        doc,
        "На каждой итерации разведчики формируют набор участков. После оценки целевой функции участки "
        "сортируются по качеству. Для элитных участков выделяется больше рабочих пчел, чем для просто "
        "перспективных, поэтому локальный поиск в лучших областях выполняется интенсивнее."
    )
    add_paragraph(doc, "Основные параметры алгоритма:")
    add_paragraph(doc, "Пчелы-разведчики — число начальных точек поиска.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, "Элитные и перспективные участки — число лучших областей, вокруг которых ведется локальный поиск.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, "Пчелы на участке — число новых кандидатов, создаваемых в окрестности найденной области.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, "Радиус участка — размер локальной области поиска вокруг найденной точки.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, "Коэффициент уменьшения радиуса — регулирует переход от грубого поиска к уточнению решения.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_blank(doc, 1)

    add_heading_like(doc, "ВЫПОЛНЕНИЕ ШАГОВ АЛГОРИТМА")
    add_blank(doc, 1)
    add_paragraph(doc, "Класс BeesAlgorithm")
    add_paragraph(
        doc,
        "Класс реализует основную механику алгоритма: инициализацию разведчиков, выделение элитных и "
        "перспективных участков, локальный поиск рабочих пчел и контроль стагнации."
    )
    add_code_line(doc, "class BeesAlgorithm:")
    add_code_line(doc, "    def _random_position(self) -> np.ndarray:")
    add_code_line(doc, "        return np.random.uniform(self.bounds[:, 0], self.bounds[:, 1], size=self.dimension)")
    add_code_line(doc, "    ")
    add_code_line(doc, "    def _random_in_patch(self, center: np.ndarray, radius: float) -> np.ndarray:")
    add_code_line(doc, "        position = center + np.random.uniform(-radius, radius, size=self.dimension)")
    add_code_line(doc, "        return np.clip(position, self.bounds[:, 0], self.bounds[:, 1])")
    add_code_line(doc, "    ")
    add_code_line(doc, "    def solve(self, real_time_callback=None):")
    add_code_line(doc, "        # сортировка участков, локальный поиск и уменьшение радиуса")
    add_code_line(doc, "        ...")
    add_blank(doc, 1)
    add_paragraph(doc, "Класс BeesAlgorithmWorker")
    add_paragraph(
        doc,
        "Для режима реального времени используется отдельный поток выполнения. "
        "Это позволяет пошагово обновлять состояние популяции и при этом не блокировать интерфейс PyQt6."
    )
    add_paragraph(doc, "Класс Lab5Widget")
    add_paragraph(
        doc,
        "Графический виджет отвечает за ввод параметров, выбор тестовой функции, запуск пакетного "
        "или realtime-режима, построение поверхности целевой функции в VTK и отображение текущего "
        "положения пчел и лучшей найденной точки."
    )
    add_blank(doc, 2)

    add_heading_like(doc, "ПРОГРАММНЫЙ ИНТЕРФЕЙС")
    add_blank(doc, 1)
    add_paragraph(
        doc,
        "Интерфейс лабораторной работы содержит выпадающий список выбора функции, поля для задания "
        "параметров пчелиного роя, переключатель realtime-режима, регулятор скорости анимации и "
        "ползунок для просмотра истории итераций."
    )
    add_paragraph(
        doc,
        "После запуска строится поверхность целевой функции, на ней отображаются позиции пчел текущей "
        "итерации, а найденный лучший минимум выделяется отдельной сферой."
    )
    add_placeholder_figure(doc, 1, "Начальный программный интерфейс")
    add_paragraph(
        doc,
        "На следующем рисунке должен быть показан процесс работы алгоритма в режиме реального времени, "
        "когда пчелы постепенно концентрируются в окрестности лучшего найденного решения."
    )
    add_placeholder_figure(doc, 2, "Визуализация работы алгоритма пчелиного роя")
    add_paragraph(
        doc,
        "В текстовой области результатов выводятся координаты лучшей точки, значение функции и "
        "основные параметры колонии."
    )
    add_placeholder_figure(doc, 3, "Результат оптимизации")
    add_blank(doc, 1)
    add_paragraph(
        doc,
        "Вывод: в ходе выполнения лабораторной работы были изучены теоретические основы алгоритма "
        "пчелиного роя, реализованы механизмы глобального и локального поиска на языке Python с "
        "использованием NumPy, PyQt6 и VTK. Разработано приложение, позволяющее визуализировать "
        "поведение пчел в пространстве поиска и наблюдать влияние параметров колонии, радиуса участка "
        "и критерия стагнации на сходимость метода.",
        bold_prefix="Вывод:",
    )
    return doc


def build_lab6():
    doc = setup_document()
    add_cover_page(doc, 6, "ИММУННАЯ СЕТЬ")

    add_paragraph(
        doc,
        "Цель: изучить применение искусственных иммунных систем для задач глобальной оптимизации, "
        "реализовать алгоритм иммунной сети на основе клональной селекции, гипермутации и супрессии, "
        "а также исследовать его работу на тестовых функциях двух переменных.",
        bold_prefix="Цель:",
    )
    add_blank(doc, 1)

    add_heading_like(doc, "ТЕОРИТИЧЕСКАЯ ЧАСТЬ")
    add_blank(doc, 1)
    add_paragraph(
        doc,
        "Искусственные иммунные системы представляют собой класс алгоритмов, вдохновленных механизмами "
        "биологического иммунитета. В оптимизационной постановке возможные решения интерпретируются "
        "как антитела, а качество решения определяется степенью их аффинности к задаче, то есть значением "
        "целевой функции."
    )
    add_paragraph(
        doc,
        "В иммунной сети лучшие антитела отбираются, клонируются и подвергаются гипермутации. "
        "Чем выше качество текущего решения, тем точнее выполняется локальное уточнение его окрестности. "
        "Дополнительно применяется супрессия, устраняющая слишком близкие между собой решения, чтобы "
        "поддерживать разнообразие популяции."
    )
    add_paragraph(doc, "Основные этапы алгоритма:")
    add_paragraph(doc, "Инициализация популяции антител случайными точками в области поиска.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, "Оценка аффинности и отбор лучших антител.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, "Клонирование и гипермутация выбранных решений.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, "Супрессия близких антител для сохранения разнообразия.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, "Добавление новых случайных антител для глобального поиска.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_blank(doc, 1)

    add_heading_like(doc, "ВЫПОЛНЕНИЕ ШАГОВ АЛГОРИТМА")
    add_blank(doc, 1)
    add_paragraph(doc, "Класс ImmuneNetworkAlgorithm")
    add_paragraph(
        doc,
        "Класс реализует клональную селекцию и механизмы обновления популяции. "
        "На каждой итерации лучшие антитела клонируются, затем модифицируются гипермутацией, "
        "после чего из полученного множества отбираются наиболее разнообразные и качественные решения."
    )
    add_code_line(doc, "class ImmuneNetworkAlgorithm:")
    add_code_line(doc, "    def _hypermutate(self, antibody: np.ndarray, score: float) -> np.ndarray:")
    add_code_line(doc, "        normalized_score = 1.0 / (1.0 + max(score, 0.0))")
    add_code_line(doc, "        scale = self.mutation_scale * (1.2 - normalized_score)")
    add_code_line(doc, "        delta = np.random.normal(0, scale, size=self.dimension)")
    add_code_line(doc, "        return np.clip(antibody + delta, self.bounds[:, 0], self.bounds[:, 1])")
    add_code_line(doc, "    ")
    add_code_line(doc, "    def _suppress_population(self, population, values):")
    add_code_line(doc, "        # удаление слишком близких решений")
    add_code_line(doc, "        ...")
    add_blank(doc, 1)
    add_paragraph(doc, "Класс ImmuneNetworkWorker")
    add_paragraph(
        doc,
        "Для визуализации по итерациям создан рабочий поток, который передает в интерфейс текущее "
        "состояние популяции без блокировки главного окна приложения."
    )
    add_paragraph(doc, "Класс Lab6Widget")
    add_paragraph(
        doc,
        "Виджет лабораторной работы позволяет выбирать тестовую функцию, число антител, размер "
        "клонального набора, силу гипермутации, радиус супрессии и параметры realtime-визуализации."
    )
    add_blank(doc, 2)

    add_heading_like(doc, "ПРОГРАММНЫЙ ИНТЕРФЕЙС")
    add_blank(doc, 1)
    add_paragraph(
        doc,
        "Программный интерфейс содержит элементы управления для задания параметров иммунной сети: "
        "размер популяции, количество отбираемых антител, число клонов, масштаб гипермутации, "
        "радиус супрессии и количество новых антител, добавляемых на каждой итерации."
    )
    add_paragraph(
        doc,
        "В realtime-режиме пользователь может наблюдать перестройку популяции антител на поверхности "
        "целевой функции и отслеживать перемещение лучшего решения."
    )
    add_placeholder_figure(doc, 1, "Начальный программный интерфейс")
    add_paragraph(
        doc,
        "На следующем рисунке должен быть показан процесс эволюции иммунной сети, когда популяция "
        "постепенно концентрируется в перспективной области пространства поиска."
    )
    add_placeholder_figure(doc, 2, "Визуализация работы иммунной сети")
    add_paragraph(
        doc,
        "В текстовой области выводятся итоговые координаты минимума, значение функции и использованные параметры сети."
    )
    add_placeholder_figure(doc, 3, "Результат оптимизации")
    add_blank(doc, 1)
    add_paragraph(
        doc,
        "Вывод: в ходе выполнения лабораторной работы были изучены принципы искусственных иммунных "
        "систем и реализован алгоритм иммунной сети для глобальной оптимизации. Программная реализация "
        "на Python с использованием NumPy, PyQt6 и VTK позволила исследовать влияние клональной селекции, "
        "гипермутации и супрессии на качество поиска и устойчивость алгоритма при работе с многомодальными "
        "функциями.",
        bold_prefix="Вывод:",
    )
    return doc


def build_lab7():
    doc = setup_document()
    add_cover_page(doc, 7, "БАКТЕРИАЛЬНАЯ ОПТИМИЗАЦИЯ")

    add_paragraph(
        doc,
        "Цель: изучить принципы бактериальной оптимизации как биоинспирированного метода глобального "
        "поиска, реализовать процессы chemotaxis, swim, reproduction и elimination-dispersal, "
        "а также исследовать сходимость алгоритма при минимизации тестовых функций двух переменных.",
        bold_prefix="Цель:",
    )
    add_blank(doc, 1)

    add_heading_like(doc, "ТЕОРИТИЧЕСКАЯ ЧАСТЬ")
    add_blank(doc, 1)
    add_paragraph(
        doc,
        "Бактериальная оптимизация моделирует поведение бактерий при поиске благоприятной среды. "
        "Каждая бактерия перемещается по пространству поиска, выполняя случайный поворот (tumble) "
        "и продолжая движение в успешном направлении (swim), если новое положение улучшает значение "
        "целевой функции."
    )
    add_paragraph(
        doc,
        "После серии шагов chemotaxis проводится размножение: более успешные бактерии сохраняются, "
        "а менее успешные заменяются копиями лучших. Для предотвращения преждевременной сходимости "
        "часть бактерий периодически рассеивается случайным образом по области поиска."
    )
    add_paragraph(doc, "Основные компоненты алгоритма:")
    add_paragraph(doc, "Chemotaxis — локальное перемещение бактерий по направлению поиска.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, "Swim — продолжение движения в направлении улучшения.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, "Reproduction — сохранение лучших бактерий и дублирование наиболее успешных.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, "Elimination-dispersal — случайное рассеивание части популяции.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_blank(doc, 1)

    add_heading_like(doc, "ВЫПОЛНЕНИЕ ШАГОВ АЛГОРИТМА")
    add_blank(doc, 1)
    add_paragraph(doc, "Класс BacterialForagingOptimization")
    add_paragraph(
        doc,
        "Класс реализует основной цикл бактериальной оптимизации. Внутри него последовательно "
        "выполняются шаги chemotaxis, накопление меры здоровья бактерий, размножение и этапы "
        "рассеивания популяции."
    )
    add_code_line(doc, "class BacterialForagingOptimization:")
    add_code_line(doc, "    def _tumble(self, position: np.ndarray) -> np.ndarray:")
    add_code_line(doc, "        direction = np.random.normal(size=self.dimension)")
    add_code_line(doc, "        normalized_direction = direction / np.linalg.norm(direction)")
    add_code_line(doc, "        step = normalized_direction * self.step_size * (self.bounds[:, 1] - self.bounds[:, 0])")
    add_code_line(doc, "        return np.clip(position + step, self.bounds[:, 0], self.bounds[:, 1])")
    add_code_line(doc, "    ")
    add_code_line(doc, "    def solve(self, real_time_callback=None):")
    add_code_line(doc, "        # chemotaxis, swim, reproduction, elimination-dispersal")
    add_code_line(doc, "        ...")
    add_blank(doc, 1)
    add_paragraph(doc, "Класс BacterialOptimizationWorker")
    add_paragraph(
        doc,
        "Для режима реального времени реализован отдельный поток, который передает интерфейсу "
        "текущие положения бактерий и значение лучшего найденного решения."
    )
    add_paragraph(doc, "Класс Lab7Widget")
    add_paragraph(
        doc,
        "Графический виджет предоставляет пользователю параметры размера популяции, числа шагов "
        "chemotaxis, длины swim, числа циклов размножения и рассеивания, а также настройки скорости анимации."
    )
    add_blank(doc, 2)

    add_heading_like(doc, "ПРОГРАММНЫЙ ИНТЕРФЕЙС")
    add_blank(doc, 1)
    add_paragraph(
        doc,
        "Интерфейс лабораторной работы включает выбор тестовой функции, поля ввода параметров "
        "бактериальной оптимизации, режим реального времени и ползунок для просмотра истории итераций."
    )
    add_paragraph(
        doc,
        "Во время работы алгоритма бактерии отображаются как отдельные точки на поверхности целевой "
        "функции, а лучший найденный минимум выделяется отдельной сферой."
    )
    add_placeholder_figure(doc, 1, "Начальный программный интерфейс")
    add_paragraph(
        doc,
        "На следующем рисунке должен быть показан пример распределения бактерий в процессе поиска "
        "и их постепенное смещение в область лучшего минимума."
    )
    add_placeholder_figure(doc, 2, "Визуализация работы бактериальной оптимизации")
    add_paragraph(
        doc,
        "В итоговой области результатов отображаются координаты найденного минимума, значение функции "
        "и значения параметров алгоритма."
    )
    add_placeholder_figure(doc, 3, "Результат оптимизации")
    add_blank(doc, 1)
    add_paragraph(
        doc,
        "Вывод: в ходе выполнения лабораторной работы были изучены основные механизмы бактериальной "
        "оптимизации и реализован алгоритм, включающий chemotaxis, swim, reproduction и "
        "elimination-dispersal. Создано приложение на Python с использованием NumPy, PyQt6 и VTK, "
        "позволяющее визуализировать процесс поиска и анализировать влияние параметров движения и "
        "рассеивания бактерий на скорость и устойчивость сходимости.",
        bold_prefix="Вывод:",
    )
    return doc


def build_lab8():
    doc = setup_document()
    add_cover_page(doc, 8, "ГИБРИДНЫЙ АЛГОРИТМ ОПТИМИЗАЦИИ")

    add_paragraph(
        doc,
        "Цель: разработать гибридный алгоритм оптимизации функции двух переменных, объединив два ранее "
        "реализованных метода, исследовать их совместную работу и оценить эффективность гибридного "
        "подхода при решении задач глобальной оптимизации.",
        bold_prefix="Цель:",
    )
    add_blank(doc, 1)

    add_heading_like(doc, "ТЕОРИТИЧЕСКАЯ ЧАСТЬ")
    add_blank(doc, 1)
    add_paragraph(
        doc,
        "Гибридные алгоритмы оптимизации объединяют сильные стороны нескольких методов поиска. "
        "Такой подход позволяет компенсировать недостатки каждого отдельного алгоритма и улучшить "
        "качество получаемых решений. В данной лабораторной работе в качестве гибридизируемых методов "
        "использованы алгоритм роя частиц и алгоритм пчелиного роя."
    )
    add_paragraph(
        doc,
        "Алгоритм роя частиц хорошо подходит для быстрого глобального исследования пространства поиска. "
        "Он эффективно находит перспективные области за счет обмена информацией между частицами. "
        "Алгоритм пчелиного роя, в свою очередь, удобен для локального уточнения решения, поскольку "
        "выполняет детальное исследование окрестностей лучших найденных точек."
    )
    add_paragraph(
        doc,
        "В гибридной схеме сначала выполняется PSO-фаза, на которой определяется набор перспективных "
        "областей. Затем лучшие найденные частицы передаются во вторую фазу, где пчелиный алгоритм "
        "осуществляет локальный поиск с уменьшающимся радиусом. Благодаря этому гибридный алгоритм "
        "сочетает глобальный обзор пространства и точное локальное уточнение."
    )
    add_blank(doc, 1)

    add_heading_like(doc, "ПРАКТИЧЕСКАЯ ЧАСТЬ")
    add_blank(doc, 1)
    add_paragraph(
        doc,
        "В рамках практической части был разработан гибридный алгоритм PSO + Bees. На первой фазе "
        "используется алгоритм роя частиц, который запускается на выбранной тестовой функции и выполняет "
        "поиск перспективных областей. История движения частиц сохраняется для последующей визуализации."
    )
    add_paragraph(
        doc,
        "После завершения первой фазы лучшие позиции частиц используются как стартовые точки для "
        "пчелиного алгоритма. На второй фазе вокруг этих точек организуется локальный поиск: "
        "выделяются элитные и перспективные участки, в их окрестностях генерируются новые кандидаты, "
        "а радиус поиска постепенно уменьшается."
    )
    add_paragraph(
        doc,
        "Программная реализация выполнена на языке Python с использованием библиотек NumPy, PyQt6 и VTK. "
        "В приложении предусмотрен выбор тестовой функции, настройка параметров обеих фаз гибридного "
        "алгоритма, а также запуск в обычном и пошаговом режиме."
    )
    add_blank(doc, 1)

    add_heading_like(doc, "ВЫПОЛНЕНИЕ ШАГОВ АЛГОРИТМА")
    add_blank(doc, 1)
    add_paragraph(doc, "Класс HybridPSOBeesAlgorithm")
    add_paragraph(
        doc,
        "Класс содержит две логические части. Первая часть реализует PSO-фазу: инициализацию частиц, "
        "обновление скоростей и позиций, а также поиск глобально лучшего решения. Вторая часть реализует "
        "Bees-фазу, в которой выполняется локальный поиск вокруг лучших позиций, полученных после работы PSO."
    )
    add_code_line(doc, "class HybridPSOBeesAlgorithm:")
    add_code_line(doc, "    def solve(self, real_time_callback=None):")
    add_code_line(doc, "        # Phase 1: PSO")
    add_code_line(doc, "        # обновление скоростей, позиций и выбор лучших частиц")
    add_code_line(doc, "        ...")
    add_code_line(doc, "        # Phase 2: Bees")
    add_code_line(doc, "        # локальный поиск вокруг лучших найденных областей")
    add_code_line(doc, "        ...")
    add_blank(doc, 1)
    add_paragraph(doc, "Класс HybridPSOBeesWorker")
    add_paragraph(
        doc,
        "Для режима реального времени используется отдельный поток выполнения, который передает "
        "в графический интерфейс текущую фазу алгоритма, положение агентов и лучшее найденное решение."
    )
    add_paragraph(doc, "Класс Lab8Widget")
    add_paragraph(
        doc,
        "Графический виджет лабораторной работы предоставляет пользователю параметры обеих фаз: "
        "размер роя и число итераций PSO, а также параметры пчелиного роя для локального уточнения. "
        "На графике разными цветами отображаются агенты первой и второй фазы."
    )
    add_blank(doc, 2)

    add_heading_like(doc, "ПРОГРАММНЫЙ ИНТЕРФЕЙС")
    add_blank(doc, 1)
    add_paragraph(
        doc,
        "Интерфейс лабораторной работы содержит список выбора тестовой функции, отдельные группы "
        "параметров для PSO и Bees, переключатель режима реального времени, регулятор скорости "
        "анимации и ползунок для просмотра истории итераций."
    )
    add_paragraph(
        doc,
        "После запуска алгоритма на поверхности целевой функции сначала отображается фаза PSO, затем "
        "фаза локального уточнения пчелиным роем. Лучший найденный минимум выделяется отдельной сферой."
    )
    add_placeholder_figure(doc, 1, "Начальный программный интерфейс")
    add_paragraph(
        doc,
        "На следующем рисунке должен быть показан этап глобального поиска роем частиц, когда агенты "
        "исследуют пространство поиска и определяют перспективные области."
    )
    add_placeholder_figure(doc, 2, "Фаза глобального поиска PSO")
    add_paragraph(
        doc,
        "Следующий рисунок должен иллюстрировать работу пчелиного роя на второй фазе, где выполняется "
        "локальное уточнение найденных решений."
    )
    add_placeholder_figure(doc, 3, "Фаза локального уточнения Bees")
    add_paragraph(
        doc,
        "В итоговой текстовой области выводятся координаты найденного минимума, значение функции "
        "и параметры обеих фаз гибридного алгоритма."
    )
    add_placeholder_figure(doc, 4, "Результат работы гибридного алгоритма")
    add_blank(doc, 1)
    add_paragraph(
        doc,
        "Вывод: в ходе лабораторной работы был разработан гибридный алгоритм оптимизации, объединяющий "
        "алгоритм роя частиц и алгоритм пчелиного роя. Реализация показала, что двухэтапный подход "
        "позволяет сначала быстро выделить перспективные области поиска, а затем более точно уточнить "
        "решение. Использование Python, NumPy, PyQt6 и VTK позволило реализовать как вычислительную часть, "
        "так и наглядную визуализацию процесса оптимизации.",
        bold_prefix="Вывод:",
    )
    return doc


def main():
    reports = {
        DESKTOP / "ЛР-5 Оптимизация.docx": build_lab5(),
        DESKTOP / "ЛР-6 Оптимизация.docx": build_lab6(),
        DESKTOP / "ЛР-7 Оптимизация.docx": build_lab7(),
        DESKTOP / "ЛР-8 Оптимизация.docx": build_lab8(),
    }
    for path, doc in reports.items():
        doc.save(path)
        print(f"Saved: {path}")


if __name__ == "__main__":
    main()
